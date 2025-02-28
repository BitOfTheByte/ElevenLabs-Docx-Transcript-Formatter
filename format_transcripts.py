#!/usr/bin/env python3
import os
import re
from docx import Document
from docx.shared import RGBColor

def process_edited_files(directory='.'):
    # Look for all files with prefix "edited_" and .docx extension
    edited_files = [f for f in os.listdir(directory) if f.startswith("edited_") and f.endswith(".docx")]
    
    # Predefined color list for unique speakers
    color_list = [
        RGBColor(0, 128, 0),     # Green
        RGBColor(0, 0, 255),     # Blue
        RGBColor(255, 165, 0),   # Orange
        RGBColor(128, 0, 128),   # Purple
        RGBColor(0, 128, 128),   # Teal
        RGBColor(128, 128, 0),   # Olive
        RGBColor(255, 0, 255)    # Magenta
    ]
    
    for edited_file in edited_files:
        doc_path = os.path.join(directory, edited_file)
        doc = Document(doc_path)
        new_doc = Document()
        
        speaker_color_map = {}  # Maps speaker name to its RGBColor
        color_index = 0         # Index to track next available color
        
        paragraphs = doc.paragraphs
        i = 0
        while i < len(paragraphs):
            header = paragraphs[i].text.strip()
            # Expect header line format: ... - [SpeakerName]
            match = re.match(r'.*-\s*\[([^\]]+)\]\s*$', header)
            if match:
                speaker = match.group(1)
                # Attempt to get the dialogue text from the next paragraph
                dialogue = ""
                if i + 1 < len(paragraphs):
                    dialogue = paragraphs[i + 1].text.strip()
                i += 2  # Skip header and dialogue paragraphs
                
                # Assign a unique color for the speaker if not already done
                if speaker not in speaker_color_map:
                    speaker_color_map[speaker] = color_list[color_index % len(color_list)]
                    color_index += 1
                speaker_color = speaker_color_map[speaker]
                
                # Create a new paragraph in the formatted document:
                # Format: [SpeakerName]: Dialogue text
                new_para = new_doc.add_paragraph()
                speaker_run = new_para.add_run(f'[{speaker}]: ')
                speaker_run.font.color.rgb = speaker_color
                new_para.add_run(dialogue)
            else:
                # If paragraph does not match expected header, include it if not empty.
                if header:
                    new_doc.add_paragraph(header)
                i += 1
        
        # Save the newly formatted document with a new prefix
        new_filename = os.path.join(directory, f"formatted_{edited_file}")
        new_doc.save(new_filename)
        print(f"Processed {edited_file} and saved as {new_filename}")

if __name__ == '__main__':
    process_edited_files()
